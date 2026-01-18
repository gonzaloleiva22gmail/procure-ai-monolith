import pandas as pd
from pypdf import PdfReader
from docx import Document
import os
from pathlib import Path

def read_excel(filepath):
    try:
        df = pd.read_excel(filepath)
        return df.to_csv(index=False)
    except Exception as e:
        return f"Error reading Excel {filepath}: {str(e)}"

def read_pdf(filepath):
    try:
        reader = PdfReader(filepath)
        text = ""
        for page in reader.pages:
            text += page.extract_text() + "\n"
        return text
    except Exception as e:
        return f"Error reading PDF {filepath}: {str(e)}"

def read_word(filepath):
    try:
        doc = Document(filepath)
        full_text = []
        # Paragraphs
        for para in doc.paragraphs:
            if para.text.strip():
                full_text.append(para.text)
        # Tables
        for table in doc.tables:
            for row in table.rows:
                row_text = [cell.text for cell in row.cells]
                full_text.append(" | ".join(row_text))
        
        return "\n".join(full_text)
    except Exception as e:
        return f"Error reading Word {filepath}: {str(e)}"

def read_text(filepath):
    try:
        with open(filepath, 'r', encoding='utf-8') as f:
            return f.read()
    except Exception as e:
        return f"Error reading Text {filepath}: {str(e)}"

def get_knowledge_base_content(base_path):
    """Recursively walks through base_path and extracts text."""
    context_text = ""
    
    if not os.path.exists(base_path):
        return "No knowledge base found."

    for root, dirs, files in os.walk(base_path):
        for file in files:
            file_path = os.path.join(root, file)
            ext = file.split('.')[-1].lower()
            
            content = ""
            if ext in ['xlsx', 'xls']:
                content = read_excel(file_path)
            elif ext == 'pdf':
                content = read_pdf(file_path)
            elif ext in ['docx', 'doc']:
                content = read_word(file_path)
            elif ext == 'txt':
                content = read_text(file_path)
            
            if content:
                context_text += f"\n--- File: {file} ---\n{content}\n"
    
    return context_text
