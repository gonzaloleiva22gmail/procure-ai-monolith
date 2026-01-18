import re
import json
from docx import Document

def analyze_document(file_path: str):
    """
    Robustly scans a DOCX file for variables {{ v... }} and {{ V... }}.
    Captures full row context for table-based variables.
    """
    doc = Document(file_path)
    variables = {}
    
    # Regex to find {{ v1 }}, {{ V2 }}, etc. 
    # The flag re.IGNORECASE handles the v vs V issue.
    var_pattern = re.compile(r"\{\{\s*(v\d+)\s*\}\}", re.IGNORECASE)

    # Helper to clean text
    def clean_text(text):
        return text.replace('\xa0', ' ').strip()

    # 1. SCAN TABLES (Primary Source for Forms)
    for table in doc.tables:
        for row in table.rows:
            # Capture the entire row text as context
            row_cells = [clean_text(cell.text) for cell in row.cells if cell.text.strip()]
            full_row_context = " | ".join(row_cells)
            
            # Scan each cell
            for cell in row.cells:
                text = clean_text(cell.text)
                matches = var_pattern.findall(text)
                for var_raw in matches:
                    # Normalize to lowercase (e.g., 'V2' -> 'v2') to ensure consistency
                    var_id = var_raw.lower()
                    
                    # Store with rich context
                    variables[var_id] = {
                        "id": var_id,
                        "original_tag": var_raw,
                        "context": full_row_context,
                        "type": "table_row"
                    }

    # 2. SCAN PARAGRAPHS (Fallback)
    for para in doc.paragraphs:
        text = clean_text(para.text)
        matches = var_pattern.findall(text)
        for var_raw in matches:
            var_id = var_raw.lower()
            if var_id not in variables:
                variables[var_id] = {
                    "id": var_id,
                    "original_tag": var_raw,
                    "context": text,
                    "type": "paragraph"
                }

    # Sort by number (v1, v2, v10...)
    def get_v_num(v_dict):
        try:
            return int(re.search(r'\d+', v_dict['id']).group())
        except:
            return 9999

    sorted_vars = sorted(variables.values(), key=get_v_num)
    return sorted_vars

if __name__ == "__main__":
    import sys
    if len(sys.argv) > 1:
        print(json.dumps(analyze_document(sys.argv[1]), indent=2))
