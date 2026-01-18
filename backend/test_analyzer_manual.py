import os
import docx
import json
from unittest.mock import MagicMock, patch
import backend.analyzer

def create_dummy_template(filename):
    """Creates a dummy DOCX template with variables in paragraphs and tables."""
    doc = docx.Document()
    
    # Add paragraph with variable
    doc.add_paragraph("This is a contract for {{ v1 }} made on {{ v2 }}.")
    
    # Add table with variable
    table = doc.add_table(rows=2, cols=3)
    
    # Header row
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = "Item"
    hdr_cells[1].text = "Description"
    hdr_cells[2].text = "Cost"
    
    # Data row
    row_cells = table.rows[1].cells
    row_cells[0].text = "Service A"
    row_cells[1].text = "Details: {{ v3 }}"
    row_cells[2].text = "$100"
    
    os.makedirs(os.path.join("backend", "templates"), exist_ok=True)
    doc.save(os.path.join("backend", "templates", filename))
    print(f"Created dummy template: {filename}")

def run_verification():
    test_file = "test_template_1.docx"
    create_dummy_template(test_file)
    
    # Patch the OpenAI client in backend.analyzer
    with patch("backend.analyzer.client") as mock_client:
        # Setup mock response
        mock_response = MagicMock()
        mock_response.choices[0].message.content = json.dumps({
            "v1": {"label": "Client", "type": "text"},
            "v2": {"label": "Date", "type": "date"},
            "v3": {"label": "Desc", "type": "text"}
        })
        mock_client.chat.completions.create.return_value = mock_response

        print("\nRunning analyze_document with mocked API...")
        result = backend.analyzer.analyze_document(test_file)
        
        # Verify result
        print("Result:", json.dumps(result, indent=2))
        
        # Verify Prompt Content to ensure extraction worked
        # Get the call arguments
        call_args = mock_client.chat.completions.create.call_args
        if not call_args:
            print("ERROR: API was not called!")
            return

        kwargs = call_args.kwargs
        messages = kwargs.get('messages', [])
        user_message = next((m['content'] for m in messages if m['role'] == 'user'), "")
        
        print("\nVerifying Prompt Content:")
        
        # Check v1 context (paragraph)
        if "v1" in user_message and "This is a contract for" in user_message:
            print("[PASS] v1 extracted with context.")
        else:
            print("[FAIL] v1 context missing.")

        # Check v3 context (table)
        # We expect: Header: 'Description', Row Label: 'Service A'
        if "v3" in user_message:
            if "Header: 'Description'" in user_message and "Row Label: 'Service A'" in user_message:
                print("[PASS] v3 extracted with Table Metadata (Header + Row Label).")
            else:
                print(f"[FAIL] v3 found but table metadata missing. Prompt snippet: {user_message[user_message.find('v3'):][:200]}")
        else:
            print("[FAIL] v3 missing from prompt.")

if __name__ == "__main__":
    run_verification()
