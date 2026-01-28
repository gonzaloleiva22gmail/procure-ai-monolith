print("\n\n!!! I AM THE NEW VERSION - IF YOU SEE THIS, THE CODE IS LOADED !!!\n\n")

import os
import sys
try:
    from dotenv import load_dotenv
    # 1. LOAD ENV
    current_dir = os.path.dirname(os.path.abspath(__file__))
    env_path = os.path.join(current_dir, '..', '.env')
    load_dotenv(dotenv_path=env_path)
except ImportError:
    print("Warning: python-dotenv not found. Assuming environment variables are set by the platform.")

# DEBUG: Print environment info
# DEBUG: Print environment info
import os
print(f"DEBUG: Starting Backend...")
print(f"DEBUG: Current Directory: {os.getcwd()}")
print(f"DEBUG: PORT env var: {os.environ.get('PORT', 'Not Set')}")
print(f"DEBUG: XAI_API_KEY present: {'Yes' if os.environ.get('XAI_API_KEY') else 'No'}")

# DEBUG: List static files
static_path = os.path.join(os.path.dirname(os.path.abspath(__file__)), "static")
print(f"DEBUG: Checking static path: {static_path}")
if os.path.exists(static_path):
    print(f"DEBUG: Static folder contents: {os.listdir(static_path)}")
    assets_path = os.path.join(static_path, "assets")
    if os.path.exists(assets_path):
        print(f"DEBUG: Assets folder contents: {os.listdir(assets_path)}")
    else:
        print("DEBUG: Assets folder NOT found inside static!")
else:
    print("DEBUG: Static folder NOT found!")

from fastapi import FastAPI, HTTPException
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import FileResponse
from fastapi.staticfiles import StaticFiles
from pydantic import BaseModel
from typing import List, Dict, Optional
from openai import OpenAI
# IMPORTS (Hybrid Strategy for Local vs Production)
try:
    # Local Development (Repo Root is path)
    from backend.analyzer import analyze_document
    from backend.generator import generate_document
    from backend.file_utils import get_knowledge_base_content
except ImportError:
    # Production / Railway (Backend folder indicates root context)
    try:
        from analyzer import analyze_document
        from generator import generate_document
        from file_utils import get_knowledge_base_content
    except ImportError as e:
        print(f"CRITICAL IMPORT ERROR: {e}")
        # Re-raise to crash logs so we can debug
        raise e

app = FastAPI()

app.add_middleware(
    CORSMiddleware,
    allow_origins=[
        "http://localhost:5173",
        "http://localhost:8000",
        "https://procure-ai-monolith-production.up.railway.app",
        "*"
    ],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

XAI_API_KEY = os.getenv("XAI_API_KEY")
client = OpenAI(api_key=XAI_API_KEY or "dummy_key", base_url="https://api.x.ai/v1")

class ChatRequest(BaseModel): 
    message: str
    filename: Optional[str] = None
    history: Optional[List[Dict[str, str]]] = []
class AnalyzeRequest(BaseModel): filename: str
class GenerateRequest(BaseModel): filename: str; answers: Dict[str, str]
class DraftRequest(BaseModel): field_label: str; user_notes: str; context_files: Optional[str] = ""

# --- ENDPOINTS ---

@app.post("/template-chat")
def template_consultant_chat(request: ChatRequest):
    import time
    request_start = time.time()
    print(f"DEBUG: Endpoint /template-chat hit with message: {request.message[:50]}...")
    
    # This is the endpoint that was missing
    analysis_context = "No template selected."
    if request.filename:
        try:
            print(f"DEBUG: Analyzing file: {request.filename}")
            # FIX: Use relative path from main.py
            file_path = os.path.join(current_dir, "templates", request.filename)
            if os.path.exists(file_path):
                print(f"DEBUG: File found at {file_path}. Starting analysis...")
                res = analyze_document(file_path)
                print(f"DEBUG: Analysis complete. Result size: {len(str(res))}")
                analysis_context = f"Variables found: {str(res)}"
            else:
                print(f"DEBUG: File NOT found at {file_path}")
        except Exception as e:
            print(f"DEBUG: Analysis failed: {e}")
            pass

    system_instruction = (
        f"You are an Intelligent Document Analyst. Context: {analysis_context}. "
        "Your Execution Logic:\n"
        "1. Analyze Context: Read the list of variables/questions provided in the context.\n"
        "2. Scan Input: Review the user's provided project brief or text.\n"
        "3. Cross-Reference: For EACH variable, check if the user's text answers it. "
        "Recognize that a single source section often answers multiple variables. Map data to ALL matching variables.\n"
        "4. Extract: If YES, extract content into 'extracted_data' JSON key. "
        "**High Fidelity Rules**:\n"
        "   - Verbatim Retention: Do NOT summarize list items. If the source text contains bullet points, extract the entire list. Retain the richness.\n"
        "   - Detail Preservation: If text mentions specific metrics (e.g., '4 hours RTO'), ensure these are explicitly preserved. Do not generalize.\n"
        "5. Report: In the 'response' key, output a strict Status Report. Do NOT be conversational.\n\n"
        "Response Format (Strictly Enforce This):\n"
        "1. MAPPED VARIABLES:\n"
        "   [Variable Name]: [Snippet of extracted text]\n...\n"
        "2. GAP ANALYSIS (MISSING):\n"
        "   [Variable Name]: [Brief description of what is needed]\n...\n\n"
        "IMPORTANT: Output valid JSON with exactly two keys: 'response' and 'extracted_data'."
    )
    messages = [{"role": "system", "content": system_instruction}]
    if request.history: messages.extend(request.history)
    messages.append({"role": "user", "content": request.message})

    try:
        print("DEBUG: Sending request to AI API (grok-3)...")
        api_start = time.time()
        completion = client.chat.completions.create(
            model="grok-3", 
            messages=messages,
            temperature=0.7,
            response_format={"type": "json_object"}
        )
        api_duration = time.time() - api_start
        print(f"DEBUG: AI API response received in {api_duration:.2f}s")
        import json
        result = json.loads(completion.choices[0].message.content)
        total_duration = time.time() - request_start
        print(f"DEBUG: Total /template-chat duration: {total_duration:.2f}s")
        return result
    except Exception as e:
        print(f"DEBUG: AI API failed: {e}")
        return {"response": f"Error: {str(e)}", "extracted_data": {}}

@app.get("/templates")
def get_templates():
    # FIX: Use relative path from main.py
    templates_dir = os.path.join(current_dir, "templates")
    if not os.path.exists(templates_dir): return []
    return [{"id": f, "name": f, "size": "Unknown", "author": "System"} for f in os.listdir(templates_dir) if f.endswith(".docx") and not f.startswith("~$")]

@app.get("/contracts")
def get_contracts():
    # FIX: Use relative path from main.py
    contracts_dir = os.path.join(current_dir, "Contracts")
    if not os.path.exists(contracts_dir): return []
    return [{"id": f, "name": f, "size": "Unknown", "author": "System"} for f in os.listdir(contracts_dir) if not f.startswith("~$")]

@app.get("/policies")
def get_policies():
    """List documents in the knowledge_base/policies folder"""
    import datetime
    policies_dir = os.path.join(current_dir, "knowledge_base", "policies")
    if not os.path.exists(policies_dir):
        return []
    
    docs = []
    for f in os.listdir(policies_dir):
        if f.startswith("~$") or os.path.isdir(os.path.join(policies_dir, f)):
            continue
            
        file_path = os.path.join(policies_dir, f)
        stats = os.stat(file_path)
        
        # Format size
        size_bytes = stats.st_size
        if size_bytes < 1024:
            size_str = f"{size_bytes} B"
        elif size_bytes < 1024 * 1024:
            size_str = f"{size_bytes / 1024:.1f} KB"
        else:
            size_str = f"{size_bytes / (1024 * 1024):.1f} MB"
            
        # Format date
        mtime = datetime.datetime.fromtimestamp(stats.st_mtime)
        date_str = mtime.strftime("%b %d, %Y")
        
        docs.append({
            "id": f,
            "name": f,
            "date": date_str,
            "size": size_str,
            "author": "Policy"
        })
    return docs

@app.post("/analyze")
def analyze_template(request: AnalyzeRequest):
    try:
        # FIX: Use relative path from main.py
        return analyze_document(os.path.join(current_dir, "templates", request.filename))
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

@app.post("/generate")
def generate_doc(request: GenerateRequest):
    try:
        output_path = generate_document(request.filename, request.answers)
        return FileResponse(output_path, filename=os.path.basename(output_path))
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

@app.post("/draft")
def draft_content(request: DraftRequest):
    try:
        completion = client.chat.completions.create(
            model="grok-3",
            messages=[
                {"role": "system", "content": "You are an expert Bid Writer."},
                {"role": "user", "content": f"Draft text for {request.field_label} based on: {request.user_notes}"}
            ]
        )
        return {"draft_text": completion.choices[0].message.content}
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

@app.post("/contract-chat")
def contract_chat(request: ChatRequest):
    """Contract Assistant endpoint - handles questions about contracts"""
    import time
    request_start = time.time()
    print(f"DEBUG: Endpoint /contract-chat hit with message: {request.message[:50]}...")
    
    # Build context based on selected contract
    contract_context = "No specific contract selected."
    if request.filename:
        try:
            print(f"DEBUG: Loading contract: {request.filename}")
            contract_path = os.path.join(current_dir, "Contracts", request.filename)
            if os.path.exists(contract_path):
                # Read contract content
                if contract_path.endswith('.txt'):
                    with open(contract_path, 'r', encoding='utf-8') as f:
                        contract_content = f.read()
                elif contract_path.endswith('.docx'):
                    from docx import Document
                    doc = Document(contract_path)
                    contract_content = '\n'.join([para.text for para in doc.paragraphs])
                else:
                    contract_content = "Unsupported file format"
                
                contract_context = f"Contract: {request.filename}\n\nContent:\n{contract_content[:5000]}"  # Limit to first 5000 chars
                print(f"DEBUG: Contract loaded successfully. Length: {len(contract_content)}")
            else:
                print(f"DEBUG: Contract file NOT found at {contract_path}")
        except Exception as e:
            print(f"DEBUG: Error loading contract: {e}")
            contract_context = f"Error loading contract: {str(e)}"
    
    system_prompt = (
        "You are a Contract Assistant specialized in analyzing procurement contracts. "
        f"Context: {contract_context}\n\n"
        "Your role is to:\n"
        "1. Answer questions about the contract content\n"
        "2. Extract key terms, dates, parties, obligations\n"
        "3. Identify risks or important clauses\n"
        "4. Provide clear, concise summaries\n\n"
        "If no specific contract is loaded, provide general contract analysis guidance."
    )
    
    messages = [{"role": "system", "content": system_prompt}]
    if request.history:
        messages.extend(request.history)
    messages.append({"role": "user", "content": request.message})
    
    try:
        print("DEBUG: Sending request to AI API (grok-3)...")
        api_start = time.time()
        completion = client.chat.completions.create(
            model="grok-3",
            messages=messages,
            temperature=0.7
        )
        api_duration = time.time() - api_start
        print(f"DEBUG: AI API response received in {api_duration:.2f}s")
        
        result = {"response": completion.choices[0].message.content}
        total_duration = time.time() - request_start
        print(f"DEBUG: Total /contract-chat duration: {total_duration:.2f}s")
        return result
    except Exception as e:
        print(f"DEBUG: AI API failed: {e}")
        return {"response": f"Error: {str(e)}"}

@app.post("/policy-chat")
def policy_chat(request: ChatRequest):
    """Policy Assistant endpoint - handles questions about policy documents"""
    import time
    request_start = time.time()
    print(f"DEBUG: Endpoint /policy-chat hit with message: {request.message[:50]}...")
    
    # Build context based on selected policy
    policy_context = "No specific policy selected."
    if request.filename:
        try:
            print(f"DEBUG: Loading policy: {request.filename}")
            # Policy documents are in backend/knowledge_base/policies
            policy_path = os.path.join(current_dir, "knowledge_base", "policies", request.filename)
            if os.path.exists(policy_path):
                # Read policy content
                # Support .txt, .pdf (if we had text extraction), .docx
                # For now, reuse the same logic as contracts (txt/docx) or generic
                if policy_path.endswith('.txt'):
                    with open(policy_path, 'r', encoding='utf-8') as f:
                        policy_content = f.read()
                elif policy_path.endswith('.docx'):
                    from docx import Document
                    doc = Document(policy_path)
                    policy_content = '\n'.join([para.text for para in doc.paragraphs])
                elif policy_path.endswith('.pdf'):
                     # Basic PDF support using pypdf if available, else placeholder
                    try:
                        from pypdf import PdfReader
                        reader = PdfReader(policy_path)
                        policy_content = ""
                        for page in reader.pages:
                            policy_content += page.extract_text() + "\n"
                    except ImportError:
                        policy_content = "PDF support requires pypdf. Please install it."
                    except Exception as e:
                        policy_content = f"Error reading PDF: {str(e)}"
                else:
                    policy_content = "Unsupported file format"
                
                policy_context = f"Policy Document: {request.filename}\n\nContent:\n{policy_content[:10000]}"  # Limit to first 10000 chars
                print(f"DEBUG: Policy loaded successfully. Length: {len(policy_content)}")
            else:
                print(f"DEBUG: Policy file NOT found at {policy_path}")
        except Exception as e:
            print(f"DEBUG: Error loading policy: {e}")
            policy_context = f"Error loading policy: {str(e)}"
    
    system_prompt = (
        "You are a Policy Assistant specialized in answering questions about company policies and procedures. "
        f"Context: {policy_context}\n\n"
        "Your role is to:\n"
        "1. Answer user questions based STRICTLY on the provided policy document.\n"
        "2. If the policy does not contain the answer, state that explicitly.\n"
        "3. Provide clear, direct answers citing the relevant section if possible.\n"
        "4. maintain a professional and helpful tone."
    )
    
    messages = [{"role": "system", "content": system_prompt}]
    if request.history:
        messages.extend(request.history)
    messages.append({"role": "user", "content": request.message})
    
    try:
        print("DEBUG: Sending request to AI API (grok-3)...")
        api_start = time.time()
        completion = client.chat.completions.create(
            model="grok-3",
            messages=messages,
            temperature=0.5 # Lower temperature for more accurate policy answers
        )
        api_duration = time.time() - api_start
        print(f"DEBUG: AI API response received in {api_duration:.2f}s")
        
        result = {"response": completion.choices[0].message.content}
        total_duration = time.time() - request_start
        print(f"DEBUG: Total /policy-chat duration: {total_duration:.2f}s")
        return result
    except Exception as e:
        print(f"DEBUG: AI API failed: {e}")
        return {"response": f"Error: {str(e)}"}

@app.post("/chat")
def chat_agent(request: ChatRequest):
    # Basic Chat
    try:
        # Load Knowledge Base Context
        # FIX: Use relative path from main.py
        kb_path = os.path.join(current_dir, "knowledge_base")
        kb_context = get_knowledge_base_content(kb_path)

        system_prompt = (
            "You are a Data Analyst. Answer the user's question based strictly on the following context:\n\n"
            f"{kb_context}\n\n"
            "If the answer is not in the context, say you don't have that information."
        )

        messages = [{"role": "system", "content": system_prompt}]
        if request.history: messages.extend(request.history)
        messages.append({"role": "user", "content": request.message})

        completion = client.chat.completions.create(
             model="grok-3",
             messages=messages
        )
        return {"response": completion.choices[0].message.content}
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

@app.get("/debug")
def debug_server():
    import glob
    templates_dir = os.path.join(current_dir, "templates")
    contracts_dir = os.path.join(current_dir, "Contracts")
    static_dir = os.path.join(current_dir, "static")
    
    return {
        "current_dir": current_dir,
        "listdir_current": os.listdir(current_dir),
        "templates_path": templates_dir,
        "templates_exists": os.path.exists(templates_dir),
        "templates_content": os.listdir(templates_dir) if os.path.exists(templates_dir) else "NOT FOUND",
        "contracts_path": contracts_dir,
        "contracts_exists": os.path.exists(contracts_dir),
        "contracts_content": os.listdir(contracts_dir) if os.path.exists(contracts_dir) else "NOT FOUND",
        "env_vars": {k: v for k, v in os.environ.items() if k in ["PORT", "XAI_API_KEY", "RAILWAY_STATIC_URL"]},
        "static_exists": os.path.exists(static_dir),
        "static_content": os.listdir(static_dir) if os.path.exists(static_dir) else "NOT FOUND",
    }

# Mount the assets folder (JS/CSS)
# FIX: Use relative path from main.py
static_assets_path = os.path.join(current_dir, "static", "assets")
if os.path.exists(static_assets_path):
    app.mount("/assets", StaticFiles(directory=static_assets_path), name="assets")

# Catch-All Route for React Router (Must be the last route)
# Note: "full_path" captures everything. "/{full_path:path}" will capture empty string too if configured, 
# but FastAPI often treats root separate. We can use a catch-all that defaults to index.html.

@app.get("/{full_path:path}")
async def serve_react_app(full_path: str):
    # FIX: Use relative path from main.py
    static_file_path = os.path.join(current_dir, "static", full_path)
    
    # 1. Try to serve specific file if it exists (e.g., favicon.ico, robot.txt)
    if full_path and os.path.exists(static_file_path) and os.path.isfile(static_file_path):
        return FileResponse(static_file_path)
    
    # DEBUG: Log missed asset requests
    if full_path.startswith("assets/"):
        print(f"DEBUG: 404 for asset: {full_path} (Looked at: {static_file_path})")
    
    # 2. Otherwise/Default: serve index.html (SPA routing)
    index_path = os.path.join(current_dir, "static", "index.html")
    if os.path.exists(index_path):
        return FileResponse(index_path)
    
    return {"error": "Frontend not built. Run 'npm run build' and move dist to backend/static"}

if __name__ == "__main__":
    import uvicorn
    uvicorn.run(app, host="0.0.0.0", port=8000)